"""
MVP Plan Document Generator
Subscription Cancellation Prediction System (OTT/SaaS)
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"c:\Users\user\Downloads\Subscription Cancellation Prediction System (OTTSaaS)\bootcamp-ace-26-team-3\Minimum_Viable_Product_MVP_Plan.docx"

# ── colour palette ──────────────────────────────────────────────────────────
C_NAVY   = RGBColor(0x0F, 0x17, 0x2A)   # headings
C_BLUE   = RGBColor(0x1D, 0x4E, 0xD8)   # sub-headings
C_STEEL  = RGBColor(0x47, 0x55, 0x69)   # body
C_ACCENT = RGBColor(0x06, 0x95, 0x8C)   # callouts / H3
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def shade_cell(cell, hex_color):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

class MVPBuilder:
    """Helper class to build a clean Word Document."""
    def __init__(self):
        self.doc = Document()
        for sec in self.doc.sections:
            sec.top_margin    = Inches(1.0)
            sec.bottom_margin = Inches(1.0)
            sec.left_margin   = Inches(1.15)
            sec.right_margin  = Inches(1.15)
        s = self.doc.styles['Normal']
        s.font.name = 'Calibri'
        s.font.size = Pt(11)
        s.font.color.rgb = C_STEEL

    def cover_title(self, txt):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.name = 'Calibri Light'; r.font.size = Pt(28)
        r.bold = True; r.font.color.rgb = C_NAVY

    def cover_sub(self, txt):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(14)
        r.italic = True; r.font.color.rgb = C_STEEL

    def h1(self, txt):
        self.doc.add_paragraph()
        p = self.doc.add_paragraph()
        r = p.add_run(txt)
        r.font.name = 'Calibri Light'; r.font.size = Pt(18)
        r.bold = True; r.font.color.rgb = C_NAVY
        p.paragraph_format.space_after = Pt(4)

    def h2(self, txt):
        p = self.doc.add_paragraph()
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(13)
        r.bold = True; r.font.color.rgb = C_BLUE
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)

    def h3(self, txt):
        p = self.doc.add_paragraph()
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(11)
        r.bold = True; r.font.color.rgb = C_ACCENT

    def body(self, txt):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(11)
        r.font.color.rgb = C_STEEL

    def bullet(self, txt, level=0):
        p = self.doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(11)

    def numbered(self, txt):
        p = self.doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt)
        r.font.name = 'Calibri'; r.font.size = Pt(11)

    def note_box(self, txt):
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run("Note: " + txt)
        r.italic = True
        r.font.color.rgb = C_ACCENT

    def spacer(self, n=1):
        for _ in range(n):
            self.doc.add_paragraph()

    def page_break(self):
        self.doc.add_page_break()

    def table(self, headers, rows, col_widths=None):
        n_cols = len(headers)
        t = self.doc.add_table(rows=1 + len(rows), cols=n_cols)
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        # header row
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            shade_cell(cell, '0F172A')
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            r.bold = True; r.font.color.rgb = C_WHITE; r.font.size = Pt(10)
        # data rows
        for ri, row in enumerate(rows):
            shade = 'F8FAFC' if ri % 2 == 0 else 'FFFFFF'
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                shade_cell(cell, shade)
                cell.text = str(val)
                cell.paragraphs[0].runs[0].font.size = Pt(10)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = Inches(w)
        self.doc.add_paragraph()

    def save(self):
        self.doc.save(OUTPUT_PATH)
        print("[OK] Document saved successfully to: " + OUTPUT_PATH)

def build_mvp_plan():
    hb = MVPBuilder()

    # ─────────────────────────────────────────────────────────────────────────
    # COVER PAGE
    # ─────────────────────────────────────────────────────────────────────────
    hb.spacer(4)
    hb.cover_title("MINIMUM VIABLE PRODUCT (MVP) PLAN")
    hb.spacer()
    hb.cover_sub("Subscription Cancellation Prediction System (OTT/SaaS)")
    hb.cover_sub("Agile Product Management & Delivery Framework")
    hb.spacer(4)
    p = hb.doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata = [
        "Document Version: 1.0 — Initial Planning Baseline",
        "Author: Senior Product Owner",
        "Target Audience: Executives, Scrum Team, Technical Stakeholders",
        "Date: July 15, 2026",
    ]
    for line in metadata:
        r = p.add_run(line + "\n")
        r.font.name = 'Calibri'; r.font.size = Pt(11); r.bold = True
        r.font.color.rgb = C_NAVY
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # REVISION HISTORY
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("1. Document Control & Revision History")
    hb.table(
        ["Version", "Date", "Author", "Reviewers", "Status / Key Changes"],
        [
            ["0.1", "10 Jun 2026", "Product Owner", "Dev Team", "Initial draft; defined MVP boundaries"],
            ["0.5", "25 Jun 2026", "Product Owner", "Scrum Master, TL", "Added detailed sprint backlog allocations"],
            ["1.0", "15 Jul 2026", "Product Owner", "Business Owner, Tech Lead", "Baseline release for MVP execution audit"],
        ],
        [1.0, 1.2, 1.3, 1.5, 2.5]
    )
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # EXECUTIVE SUMMARY
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("2. Executive Summary")
    hb.body(
        "This Minimum Viable Product (MVP) Plan outlines the product strategy, roadmap, and delivery steps "
        "for the Subscription Cancellation Prediction System. Customer churn in subscription-based models (OTT "
        "and SaaS) poses a severe risk to Monthly Recurring Revenue (MRR) and increases Customer Acquisition "
        "Costs (CAC). The goal of this product is to transition customer success operations from a reactive "
        "post-cancellation process to a proactive retention workflow powered by machine learning."
    )
    hb.body(
        "To achieve rapid time-to-market and validate the predictive value of our data, the MVP establishes a "
        "lightweight three-tier architecture: a React SPA on the frontend, a FastAPI REST API on the backend, "
        "and a structured relational schema (SQLite locally, transitioning to PostgreSQL in production). "
        "Crucially, the MVP leverages a scientifically validated Logistic Regression model that achieves a "
        "ROC-AUC of 0.9463 and a Recall of 88.71%. This model was selected as the champion after a "
        "rigorous comparison of five algorithms, where more complex methods (such as CatBoost, XGBoost, and "
        "Random Forest) were rejected due to severe overfitting (achieving suspect perfect 1.0 metrics). "
        "This plan documents the delivery of the complete system across five disciplined Agile sprints."
    )
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # PRODUCT VISION & OBJECTIVES
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("3. Product Vision & Business Objectives")
    hb.h2("3.1 Product Vision")
    hb.body(
        "To empower customer retention teams with predictive intelligence, allowing them to detect "
        "at-risk subscribers, understand their individual drivers through explainable AI, and execute "
        "targeted retention interventions before the cancellation event occurs."
    )
    hb.h2("3.2 Core Business Objectives")
    hb.bullet("Proactive Risk Identification: Automatically score customer accounts based on platform activity.")
    hb.bullet("Prescriptive Actions: Generate logical retention offers (e.g., discounts, feature upgrades) based on risk tier.")
    hb.bullet("Revenue Protection: Display Monthly Revenue at Risk as the primary KPI to drive business urgency.")
    hb.bullet("Data Portability: Allow batch analysis of external cohorts (like new trial signups) via isolated bulk upload.")
    hb.bullet("Responsive Console: Ensure the analytics dashboard loads in under 300ms to maintain user engagement.")

    hb.h2("3.3 Success Metrics (KPIs)")
    hb.table(
        ["Metric Category", "Target Metric", "Strategic Importance"],
        [
            ["Model Performance", "ROC-AUC >= 0.85, Recall >= 0.80", "High recall ensures we do not miss churners."],
            ["API Latency",       "Average response time < 500ms",  "Maintains client usability and responsive charts."],
            ["Data Integrity",    "Zero bulk upload contamination", "Keeps bulk cohort tests isolated from core database."],
            ["Query Efficiency",  "Dashboard loads under 300ms",    "Required to solve database lock limitations."],
        ],
        [2.0, 2.2, 2.3]
    )
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # MVP SCOPE & ROADMAP
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("4. MVP Definition & Scope")
    hb.h2("4.1 MVP Boundary Definition")
    hb.body(
        "The MVP is defined as the smallest functional set of features capable of scoring customer churn risk, "
        "explaining that risk to non-technical agents, and suggesting a basic retention recommendation. "
        "Features that add operational complexity without directly solving the core prediction problem have "
        "been postponed to subsequent releases."
    )
    hb.h2("4.2 MoSCoW Prioritization Matrix")
    hb.table(
        ["Feature Category", "In MVP (Must Have / Should Have)", "Postponed (Could Have / Won't Have)", "Product Rationale"],
        [
            ["Authentication", "JWT (HS256) Login & SignUp", "Multi-factor (MFA), SSO", "Basic secure access is mandatory; advanced auth is postponed."],
            ["ML Engine", "Logistic Regression (Standardized)", "CatBoost (automatic swap UI)", "CatBoost overfitted (1.0 ROC-AUC); LR is robust and stable."],
            ["Explainability", "Local SHAP (Top 3 drivers)", "Interactive global dependency plots", "Agents need simple local drivers; complex plots are too technical."],
            ["Database", "SQLite locally (Postgres prod setup)", "Automatic multi-master replication", "Zero-config SQLite fits MVP timeline; replica logic is post-MVP."],
            ["Uploads", "Isolated CSV Bulk Prediction", "Direct CRM database sync integrations", "CSV upload satisfies trial analysis; CRM sync is high effort."],
            ["Reporting", "Streaming CSV, Programmatic PDF", "Interactive custom report builder", "Static PDF and CSV exports satisfy basic audit requirements."],
            ["Dashboard", "KPI cards & 11 core charts", "Real-time WebSocket metric updates", "Periodic HTTP polling is simple; WebSockets add unnecessary overhead."],
            ["Actions", "Rule-based recommendation mapping", "Reinforcement learning for offers", "Rules are legally auditable; ML-driven pricing is too risky."],
        ],
        [1.5, 1.8, 1.8, 1.4]
    )
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # SPRINT-BY-SPRINT ROADMAP
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("5. Incremental Sprint Delivery Roadmap")
    hb.body(
        "The project was executed across five structured sprints, moving from core system and database "
        "foundations to model validation, user interface development, bulk processing, and final performance "
        "tuning."
    )

    # ── SPRINT 1 ─────────────────────────────────────────────────────────────
    hb.h2("5.1 Sprint 1: System Foundation & Secure Authentication")
    hb.h3("Goal & Business Objective")
    hb.body(
        "Establish the database schema, Pydantic validation models, and secure JWT-based login/signup "
        "infrastructure. The business objective is to ensure that customer demographic data is protected "
        "and administrative roles are restricted before predicting churn."
    )
    hb.table(
        ["Feature / Task", "User Story", "Acceptance Criteria"],
        [
            ["Database Setup", "As a developer, I want to define the users and customers schema so we can persist core records.", "Tables must be initialized in SQLite; foreign keys must be active."],
            ["JWT Login", "As a manager, I want to authenticate securely using my credentials so I can access sensitive data.", "POST /auth/login returns JWT token with 120-minute expiry on success; 401 on failure."],
            ["Admin Signup", "As an administrator, I want to create new manager accounts so I can delegate retention work.", "Only authenticated 'admin' can hit signup route; passwords hashed with pbkdf2_sha256."],
        ],
        [1.8, 2.2, 2.5]
    )
    hb.body("**Dependencies:** None. This is the foundational sprint.")
    hb.body("**Risks:** Storing weak passwords. **Mitigation:** Enforced a robust password strength regex checked server-side.")
    hb.body("**Definition of Done:** 100% test coverage on auth routes; code reviewed by Team Leader; database schema committed to Git.")
    hb.body("**Sprint Review Outcome:** Successfully authenticated client requests; default admin account seeded.")

    # ── SPRINT 2 ─────────────────────────────────────────────────────────────
    hb.h2("5.2 Sprint 2: Predictive Core & Model Comparison")
    hb.h3("Goal & Business Objective")
    hb.body(
        "Build the end-to-end ML pipeline, run the multi-model comparison script, and integrate the selected "
        "champion model. The business objective is to deliver a highly accurate, generalisable model that "
        "avoids target leakage."
    )
    hb.table(
        ["Feature / Task", "User Story", "Acceptance Criteria"],
        [
            ["Leakage Check", "As a data engineer, I want to drop columns with high target correlation so the model doesn't overfit.", "check_data_leakage() must drop Customer_ID and columns with correlation >0.95."],
            ["Model Compare", "As a PO, I want to train multiple models and compare them so we select the best generaliser.", "Comparison script compares LR, RF, GB, XGBoost, and CatBoost on a stratified split."],
            ["Champion Save", "As a developer, I want to serialize the champion model and preprocessor so the API can load them.", "Save winner to model_artifacts/; write metrics to model_metrics.json."],
        ],
        [1.8, 2.2, 2.5]
    )
    hb.body("**Dependencies:** Database schema initialized (Sprint 1).")
    hb.body(
        "**Risks:** Overfitting on tree-based models (RF, CatBoost, XGBoost all scored a perfect 1.0). "
        "**Mitigation:** Disqualified all models with ROC-AUC >= 0.999 or train/test gap > 5%; selected Logistic Regression as the stable champion."
    )
    hb.body("**Definition of Done:** Comparison reports (ROC curves, confusion matrix) saved to reports/; preprocessor serialized via pickle.")
    hb.body("**Sprint Review Outcome:** Tech Lead approved the selection of Logistic Regression (ROC-AUC 0.9463, CV ROC-AUC 0.9374).")

    # ── SPRINT 3 ─────────────────────────────────────────────────────────────
    hb.h2("5.3 Sprint 3: Customer Profiles & Direct Interventions")
    hb.h3("Goal & Business Objective")
    hb.body(
        "Implement the Customer Directory, Customer Profile view, SHAP-based local explainability, and "
        "the recommendation engine. The business objective is to present actionable, explainable churn "
        "data directly to customer success agents."
    )
    hb.table(
        ["Feature / Task", "User Story", "Acceptance Criteria"],
        [
            ["Customer Directory", "As an agent, I want to search and filter customers by risk tier so I can prioritize my tasks.", "Paginated directory loading in <200ms; filter by High/Medium/Low risk."],
            ["SHAP Integration", "As an agent, I want to see why a customer is at risk so I can customize my conversation.", "Local explanation displays top 3 features with direction labels ('Increases Risk')."],
            ["Risk Bucketing", "As a manager, I want the system to suggest offers based on risk so agents act consistently.", "High Risk (>=70%) gets 'Offer Discount'; Medium Risk (30-70%) gets 'Upgrade'; Low gets 'No Action'."],
        ],
        [1.8, 2.2, 2.5]
    )
    hb.body("**Dependencies:** Model serialized and preprocessor loaded (Sprint 2).")
    hb.body(
        "**Risks:** TreeExplainer compatibility warnings due to active Logistic Regression model. "
        "**Mitigation:** Developed custom fallback in SHAPExplainer to use generic Explainer/LinearExplainer when TreeExplainer is incompatible."
    )
    hb.body("**Definition of Done:** Directory and profile pages complete; SHAP outputs saved to explainability_json column in DB.")
    hb.body("**Sprint Review Outcome:** Product Owner approved the local explanation visualization; verified simulated risk adjustments.")

    # ── SPRINT 4 ─────────────────────────────────────────────────────────────
    hb.h2("5.4 Sprint 4: Bulk Ingestion & Executive Reporting")
    hb.h3("Goal & Business Objective")
    hb.body(
        "Develop the Bulk Prediction Studio, background processing tasks, and CSV/PDF report generators. "
        "The business objective is to support cohort analysis of unverified users without contaminating the core database."
    )
    hb.table(
        ["Feature / Task", "User Story", "Acceptance Criteria"],
        [
            ["Bulk Upload", "As a manager, I want to upload a CSV of new users so I can batch-score their churn risk.", "CSV uploaded via POST /predictions/bulk; processes in background using isolated tables."],
            ["Isolated Schema", "As a developer, I want to keep bulk results separate from customers so KPI cards remain accurate.", "Bulk results written to bulk_prediction_results; zero impact on /dashboard/kpis."],
            ["PDF Export", "As an executive, I want a PDF summary of the bulk run so I can share it with key stakeholders.", "Generates ReportLab PDF with KPIs, risk distribution table, and top high-risk users."],
        ],
        [1.8, 2.2, 2.5]
    )
    hb.body("**Dependencies:** Customer directory complete (Sprint 3).")
    hb.body(
        "**Risks:** Large uploads blocking the server event loop. "
        "**Mitigation:** Implemented FastAPI BackgroundTasks and streamed large CSV outputs row-by-row using Python generators (O(1) memory)."
    )
    hb.body("**Definition of Done:** Isolated bulk tables added to models; PDF output verified as uncorrupted binary data.")
    hb.body("**Sprint Review Outcome:** Client accepted the isolated bulk studio; verified 2,000 row import completed in under 15 seconds.")

    # ── SPRINT 5 ─────────────────────────────────────────────────────────────
    hb.h2("5.5 Sprint 5: Analytics Dashboards & Performance Optimization")
    hb.h3("Goal & Business Objective")
    hb.body(
        "Build the 11 charts for the Analytics Dashboard, resolve the SQLite database lock issue, and execute "
        "final integration tests. The business objective is to deliver a highly responsive, robust, and tested MVP."
    )
    hb.table(
        ["Feature / Task", "User Story", "Acceptance Criteria"],
        [
            ["11 Analytics Charts", "As an executive, I want to view churn rates grouped by categories so I can plan campaigns.", "Charts render churn by income, device, payment mode, age, and spend distribution."],
            ["View Optimization", "As a developer, I want to optimize the database query speed so the dashboard loads under 300ms.", "v_customer_predictions rewritten from window function to CTE-based MAX ID; runs <50ms."],
            ["Quality Assurance", "As a PO, I want 100% of integration tests to pass so I can verify release readiness.", "104 pytest integration tests run and pass; database rolled back after each test."],
        ],
        [1.8, 2.2, 2.5]
    )
    hb.body("**Dependencies:** Bulk and core tables fully populated (Sprints 1–4).")
    hb.body(
        "**Risks:** SQLite serializing parallel read/write requests, causing timeouts. "
        "**Mitigation:** Replaced ROW_NUMBER() OVER query with indexed MAX ID join, delivering a 6.75× query speedup (270ms to 40ms)."
    )
    hb.body("**Definition of Done:** 104 tests pass; frontend build complete; Render blueprint configured and validated.")
    hb.body("**Sprint Review Outcome:** Dashboard load time verified under 300ms; release baseline signed off.")
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # FEATURE PRIORITIZATION JUSTIFICATION
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("6. Feature Prioritization Strategy")
    hb.body(
        "The prioritization strategy utilized the MoSCoW framework. Prioritization was driven by two factors: "
        "Business Urgency (impact on MRR) and Technical Dependencies (foundational layers first)."
    )
    hb.h2("6.1 Prioritization Matrix Rationale")
    hb.table(
        ["Feature", "Priority", "Business Justification", "Technical Justification"],
        [
            ["JWT Security", "MUST HAVE", "Required for data privacy and regulatory compliance (GDPR/CCPA).", "Must be built first to secure all subsequent API routes."],
            ["Logistic Regression", "MUST HAVE", "Ensures immediate, generalisable churn risk predictions with zero cost.", "Simplest model with robust coefficients; avoids overfitting."],
            ["CTE View Optimization", "MUST HAVE", "Slow dashboard causes users to abandon the system.", "Eliminates SQLite write lock contention from parallel API queries."],
            ["Bulk Prediction Isolation", "MUST HAVE", "Enables testing of cohort data without contaminating KPI cards.", "Avoids expensive database joins and maintains integrity."],
            ["ReportLab PDF Export", "SHOULD HAVE", "Executives require shareable summaries of cohort prediction runs.", "Generates pure Python documents with zero system dependencies."],
            ["Interactive UI model swap", "COULD HAVE", "Nice for technical administrators but of no value to customer success agents.", "Requires database migrations and model version history interface."],
            ["Automated Model retraining", "COULD HAVE", "Addresses model drift over time, which is a slow process.", "Can be handled via offline Cron scripts for the initial release."],
            ["Real-time WebSockets", "WON'T HAVE", "User profiles do not change second-by-second; HTTP polling is sufficient.", "Adds significant state management complexity on the frontend."],
        ],
        [1.5, 0.9, 2.5, 2.1]
    )
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # RISKS, ASSUMPTIONS, CONSTRAINTS
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("7. Risks, Assumptions, Constraints, & Dependencies")
    hb.h2("7.1 Key Risks & Mitigations")
    hb.table(
        ["Risk Description", "Impact", "Mitigation Strategy"],
        [
            ["Model Drift over time", "Medium", "Log model performance statistics to the database for weekly review; prepare retraining scripts."],
            ["SQLite DB lock under load", "High", "View optimization using CTE-based MAX ID join; database connection pool limits tuned."],
            ["SHAP import failures on server", "High", "SHAP import wrapped in try/except; local explanation degrades gracefully without crashing."],
            ["CORS wildcards in production", "Medium", "Update CORSMiddleware configuration to restrict origins to the frontend domain before deployment."],
        ],
        [2.2, 1.3, 3.5]
    )

    hb.h2("7.2 Critical Assumptions")
    hb.bullet("Standard Input Format: Assumes bulk uploads contain a header matching the standard 12 feature columns.")
    hb.bullet("Browser Environment: Assumes users run modern browsers supporting local storage and secure cookies.")
    hb.bullet("Data Availability: Assumes customer demographics and ticket interactions are updated daily by backend cron jobs.")

    hb.h2("7.3 Constraints")
    hb.bullet("Single Database Writer: SQLite locks the entire database for writes, limiting parallel bulk runs.")
    hb.bullet("Memory Limits: Free-tier Render hosting limits instances to 512MB RAM, restricting the size of parsed CSVs.")
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # MVP VALIDATION STRATEGY
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("8. MVP Validation & Release Readiness")
    hb.h2("8.1 Validation Strategy")
    validation_steps = [
        "1. Functional Verification: Execute the 104 integration tests to ensure API correctness and regression coverage.",
        "2. Pilot Phase: Deploy to a staging environment and onboard three customer success agents to run predictions manually.",
        "3. Shadow Execution: Run the Logistic Regression model in parallel with legacy customer tracking for two weeks.",
        "4. Feedback Review: Analyze user actions to refine the risk scoring thresholds (30/70) and discount recommendations.",
    ]
    for step in validation_steps:
        hb.body(step)

    hb.h2("8.2 Release Readiness Checklist")
    hb.table(
        ["Check Item", "Target / Requirement", "Status"],
        [
            ["Code Coverage", "Pass all 104 backend integration tests", "✅ PASSED"],
            ["Database Schema", "Ensure v_customer_predictions view uses optimised CTE", "✅ PASSED"],
            ["Security Configuration", "Password hashing active; secrets moved to env vars", "✅ PASSED"],
            ["Deployment Setup", "Render blueprint (render.yaml) configured and tested", "✅ PASSED"],
            ["Documentation", "Master Viva Handbook and MVP Plan generated", "✅ PASSED"],
        ],
        [2.5, 3.0, 1.5]
    )
    hb.page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # POST-MVP ROADMAP
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("9. Post-MVP Roadmap & Future Enhancements")
    hb.body(
        "Following the successful launch of the MVP, the product backlog will be refined to address "
        "advanced operational features and scalability improvements."
    )
    roadmap_items = [
        "Phase 2: Automated Retraining — Implement scheduled weekly training runs to prevent model drift.",
        "Phase 2: Model Swapping UI — Allow administrators to upload new models and deploy them dynamically.",
        "Phase 3: PostgreSQL Migration — Switch the database backend to PostgreSQL to resolve write-lock limitations.",
        "Phase 3: CRM Integration — Connect directly to Salesforce or HubSpot APIs to sync customer profiles automatically.",
        "Phase 4: Proactive Alerts — Integrate Twilio/SendGrid to email agents when high-risk accounts are flagged.",
    ]
    for item in roadmap_items:
        hb.bullet(item)
    hb.spacer()

    # ─────────────────────────────────────────────────────────────────────────
    # LESSONS LEARNED
    # ─────────────────────────────────────────────────────────────────────────
    hb.h1("10. Lessons Learned & Retrospective Notes")
    hb.h2("10.1 Key Technical Lessons")
    hb.bullet("Address Database Performance Early: SQLite is fine for development, but parallel queries require careful query tuning.")
    hb.bullet("Keep Models Generalisable: Perfect training accuracy (1.0) is usually a sign of leakage, not a perfect model.")
    hb.bullet("Isolate Experimental Data: Separating bulk uploads from production customers prevents KPI corruption.")

    hb.h2("10.2 Process & Team Lessons")
    hb.bullet("Automate Testing: Having 104 integration tests allowed the team to optimize queries in Sprint 5 with confidence.")
    hb.bullet("Define Clear Acceptance Criteria: Clear definitions of done kept sprints focused and prevented scope creep.")
    hb.spacer(2)

    # Save
    hb.save()

if __name__ == "__main__":
    build_mvp_plan()
